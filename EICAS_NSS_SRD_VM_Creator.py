'''
Created on Dec 18, 2017


All EICAS headers of importance contain this exact string: "_EICAS_"
Example: AIL_TEMP_EICAS_0003

@author: logans
'''

import docx
from openpyxl import Workbook
from os import walk

def main():
    SRD_document_list = []
    for (dirpath, dirnames, filenames) in walk("C:\\Users\\logans\\Documents\\Eclipse Projects\\EICAS_SRD_VM_Creator\\src\\Test_SRDs"):
        SRD_document_list.extend(filenames)
        break
    
    print("\n\n ** LIST OF SRD DOCUMENTS FOR PARSING ** \n\n") 
    for idx in range(0,len(SRD_document_list)):
        print(str(SRD_document_list[idx]))
    
    trace_matrix_wb = Workbook()
     
    print("\n\n ** SRD DOC PARSING OUTPUT ** \n\n") 
    
    count = 0
    for file_idx in range(0,len(SRD_document_list)):
        
        # load file_name with the name of the current file from the SRD folder
        file_name = str(SRD_document_list[file_idx])     
        # create a new doc object to hold the data from the currently selected SRD
        doc = docx.Document("Test_SRDs\\" + file_name)
        # debug print SRD file info
        print(file_name)
        print(str(len(doc.paragraphs)))
        
        # worksheet names can be max 31 characters, if they are longer the excel file is unreadable
        file_name = file_name[:29]
        
        # For each document in the SRD folder, create a new worksheet to load the data into:
        tm_active = trace_matrix_wb.create_sheet(str(file_name))
        
        # this for loop searches every paragraph in the doc object for the text "EICAS"
        # if it finds it, it adds it to a new row of the trace matrix
        for idx in range(0,len(doc.paragraphs)):            
            # Note: Some requirement tags had spaces around the EICAS string (like this: "_ EICAS_" or "_EICAS  _"
            # these strings were not found in the search, so req_num_no_whitespace is designed to remove any whitespace
            # before the string is checked for "_EICAS_"
            req_num_no_whitespace = str(doc.paragraphs[idx].text)
            
            while ' ' in req_num_no_whitespace:                
                req_num_no_whitespace = req_num_no_whitespace.replace(' ','')
                
            if "_EICAS_" in req_num_no_whitespace: 
                count += 1       
                data_excel_coord = "A" + str(count)
                data_excel_text_coord = "B" + str(count)
                if count == 1:
                    # add a header to the excel worksheet before parsing data
                    # it's ugly, don't make fun of me
                    tm_active[data_excel_coord] = "Requirement Number"
                    tm_active[data_excel_text_coord] = "Requirement Text"
                    tm_active["C1"] = "EICAS Section Name"
                    tm_active["D1"] = "EICAS Definition Tag"
                    tm_active["E1"] = "Notes"
                    
                    count += 1
                    
                    # now that the header is taken care of, populate the data:
                    data_excel_coord = "A" + str(count)
                    data_excel_text_coord = "B" + str(count)
                    req_num = doc.paragraphs[idx].text
                    req_text = doc.paragraphs[idx+1].text
                    tm_active[data_excel_coord] = req_num
                    tm_active[data_excel_text_coord] = req_text
                else:
                    # no need to print header info after it's been printed once
                    req_num = doc.paragraphs[idx].text
                    req_text = doc.paragraphs[idx+1].text
                    tm_active[data_excel_coord] = req_num
                    tm_active[data_excel_text_coord] = req_text
     
        count = 0
     
    # there is always a generic "Sheet" worksheet created, so delete it before we save
    delete_sheet = trace_matrix_wb.get_sheet_by_name("Sheet")
    trace_matrix_wb.remove_sheet(delete_sheet)
    
    
    trace_matrix_wb.save("EICAS_V8_to_NSS_SRD_Trace_Matrix.xlsx")
    print("EICAS_V8_to_NSS_SRD_Trace_Matrix.xlsx saved and closed")
     
    
main()
